"""
Sync Docs/Eduaitor_Complete_QA_Test_Plan.md -> Docs/Eduaitor_Complete_QA_Test_Plan.docx
using python-docx (pandoc not available on this machine).
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[2]
MD_PATH = ROOT / "Docs" / "Eduaitor_Complete_QA_Test_Plan.md"
DOCX_PATH = ROOT / "Docs" / "Eduaitor_Complete_QA_Test_Plan.docx"

STATUS_COLORS = {
    "Pass": RGBColor(0x1E, 0x7B, 0x34),
    "Fail": RGBColor(0xC0, 0x1C, 0x28),
    "Blocked": RGBColor(0xB8, 0x86, 0x0B),
    "Not run": RGBColor(0x55, 0x55, 0x55),
}

INLINE_CODE_RE = re.compile(r"`([^`]+)`")
BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")


def add_inline_runs(paragraph, text):
    """Very small markdown-inline renderer: **bold** and `code`."""
    tokens = []
    idx = 0
    combined_re = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    for m in combined_re.finditer(text):
        if m.start() > idx:
            tokens.append(("text", text[idx:m.start()]))
        token = m.group(0)
        if token.startswith("**"):
            tokens.append(("bold", token[2:-2]))
        else:
            tokens.append(("code", token[1:-1]))
        idx = m.end()
    if idx < len(text):
        tokens.append(("text", text[idx:]))
    if not tokens:
        tokens = [("text", text)]

    for kind, val in tokens:
        run = paragraph.add_run(val)
        if kind == "bold":
            run.bold = True
        elif kind == "code":
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)


def style_status_cell(cell, status_text):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(status_text)
    run.bold = True
    color = STATUS_COLORS.get(status_text)
    if color:
        run.font.color.rgb = color


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def parse_table_block(lines, start_idx):
    """lines[start_idx] is the header row; lines[start_idx+1] is the separator."""
    header_cells = [c.strip() for c in lines[start_idx].strip().strip("|").split("|")]
    i = start_idx + 2
    rows = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        row_cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        rows.append(row_cells)
        i += 1
    return header_cells, rows, i


def add_table(doc, header_cells, rows):
    ncols = len(header_cells)
    table = doc.add_table(rows=1, cols=ncols)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(header_cells):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        set_cell_shading(hdr[i], "2F5496")
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    result_col = None
    if header_cells and header_cells[-1].strip().lower() == "result":
        result_col = ncols - 1

    for row_cells in rows:
        cells = table.add_row().cells
        for i in range(ncols):
            text = row_cells[i] if i < len(row_cells) else ""
            if i == result_col and text.strip() in STATUS_COLORS:
                style_status_cell(cells[i], text.strip())
            else:
                cells[i].text = ""
                add_inline_runs(cells[i].paragraphs[0], text)
    return table


def main():
    md_text = MD_PATH.read_text(encoding="utf-8")
    lines = md_text.split("\n")

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    i = 0
    n = len(lines)
    in_code_block = False
    code_lines = []

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Fenced code block
        if stripped.startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                in_code_block = False
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                set_cell_shading  # no-op reference to avoid lint removal
            i += 1
            continue
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Table
        if stripped.startswith("|") and i + 1 < n and re.match(r"^\|[\s:-]+\|", lines[i + 1].strip()):
            header_cells, rows, next_i = parse_table_block(lines, i)
            add_table(doc, header_cells, rows)
            doc.add_paragraph("")
            i = next_i
            continue

        # Headings
        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            heading = doc.add_heading(level=min(level, 4))
            add_inline_runs(heading, text)
            i += 1
            continue

        # Horizontal rule
        if stripped == "---":
            p = doc.add_paragraph()
            p.add_run("_" * 60)
            i += 1
            continue

        # Bullet / numbered list
        m_bullet = re.match(r"^[-*]\s+(.*)$", stripped)
        m_num = re.match(r"^\d+\.\s+(.*)$", stripped)
        if m_bullet:
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, m_bullet.group(1))
            i += 1
            continue
        if m_num:
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, m_num.group(1))
            i += 1
            continue

        # Blank line
        if stripped == "":
            i += 1
            continue

        # Default paragraph
        p = doc.add_paragraph()
        add_inline_runs(p, stripped)
        i += 1

    doc.save(str(DOCX_PATH))
    print(f"Wrote {DOCX_PATH}")


if __name__ == "__main__":
    main()
